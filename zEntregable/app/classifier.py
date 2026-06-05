"""Módulo de Clasificación de Trámites Municipales."""
import io
import re
from datetime import datetime
from typing import Optional, Tuple

import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# ─── CATÁLOGO DE TRÁMITES ──────────────────────────────────────────────────────

TRAMITE_CATALOG = {

    # ── CRÍTICA (Alta) ────────────────────────────────────────────────────────
    "Acta de Matrimonio": {
        "keywords": (
            "matrimonio conyuge contrayentes esposo esposa union civil desposados bodas nupcial "
            "casamiento registro civil ceremonia boda casados contrato matrimonial conyugal marido "
            "mujer declaracion matrimonial nupcias desposorio"
        ),
        "priority": "CRÍTICA",
        "priority_level": 3,
        "office": "Oficina de Registro Civil",
        "office_icon": "⚖️",
        "base_score": 80,
        "type_bonus": 8,
        "color": "critica",
    },
    "Acta de Defunción": {
        "keywords": (
            "defuncion fallecimiento deceso muerte causa muerte occiso finado cadaver exhumacion "
            "fallecio difunto obito decedente registro muerte certificado defuncion partida defuncion "
            "acta mortuoria sepelio entierro mortal"
        ),
        "priority": "CRÍTICA",
        "priority_level": 3,
        "office": "Oficina de Registro Civil",
        "office_icon": "⚖️",
        "base_score": 80,
        "type_bonus": 10,
        "color": "critica",
    },
    "Acta de Nacimiento": {
        "keywords": (
            "nacimiento nacio recien nacido partida nacimiento alumbramiento bebe neonato parto "
            "certificado nacimiento inscripcion nacimiento registro civil natalidad nombre menor "
            "recien nacido padre madre fecha nacimiento hospital"
        ),
        "priority": "CRÍTICA",
        "priority_level": 3,
        "office": "Oficina de Registro Civil",
        "office_icon": "⚖️",
        "base_score": 80,
        "type_bonus": 8,
        "color": "critica",
    },
    "Denuncia por Violencia Familiar": {
        "keywords": (
            "violencia agresion maltrato violencia domestica denuncia golpes lesiones amenazas "
            "intimidacion abuso familia hogar violencia fisica psicologica sexual agresor victima "
            "proteccion medidas cautelares seguridad peligro hostigamiento"
        ),
        "priority": "CRÍTICA",
        "priority_level": 3,
        "office": "Oficina de Protección Familiar",
        "office_icon": "🛡️",
        "base_score": 80,
        "type_bonus": 15,
        "color": "critica",
    },
    "Declaración de Emergencia Sanitaria": {
        "keywords": (
            "emergencia epidemia brote cuarentena sanitaria pandemia contagio infeccion enfermedad "
            "alerta salud publica caso confirmado aislamiento transmision virus bacteria notificacion "
            "epidemiologica riesgo sanitario propagacion"
        ),
        "priority": "CRÍTICA",
        "priority_level": 3,
        "office": "Gerencia de Salud Municipal",
        "office_icon": "🏥",
        "base_score": 80,
        "type_bonus": 15,
        "color": "critica",
    },

    # ── IMPORTANTE (Media) ────────────────────────────────────────────────────
    "Licencia de Conducir": {
        "keywords": (
            "licencia conducir brevete categoria vehiculo automovilista chofer manejar automovil "
            "moto transporte conductor licencia manejar certificado aptitud revalidacion brevete "
            "clase licencia examen conduccion carnet"
        ),
        "priority": "IMPORTANTE",
        "priority_level": 2,
        "office": "Gerencia de Transportes",
        "office_icon": "🚗",
        "base_score": 50,
        "type_bonus": 5,
        "color": "importante",
    },
    "Permiso de Circulación Vehicular": {
        "keywords": (
            "circulacion rodante placa SOAT revision tecnica tarjeta propiedad automotor "
            "certificado vehicular permiso transito circulacion rodado habilitacion vehicular "
            "placa rodaje matricula tarjeta circulacion inspeccion"
        ),
        "priority": "IMPORTANTE",
        "priority_level": 2,
        "office": "Gerencia de Transportes",
        "office_icon": "🚗",
        "base_score": 50,
        "type_bonus": 5,
        "color": "importante",
    },
    "Licencia de Construcción": {
        "keywords": (
            "construccion edificacion obra planos habilitacion urbana cimiento material "
            "permiso edificar arquitecto estructura vivienda piso edificio conformidad obra "
            "proyecto construccion presupuesto metros cuadrados area construida licencia obra"
        ),
        "priority": "IMPORTANTE",
        "priority_level": 2,
        "office": "Gerencia de Urbanismo y Obras",
        "office_icon": "🏗️",
        "base_score": 50,
        "type_bonus": 6,
        "color": "importante",
    },
    "Licencia de Funcionamiento": {
        "keywords": (
            "funcionamiento negocio establecimiento comercial apertura local empresa tienda "
            "mercado comercio autorizacion actividad economica zona comercial ruc registro "
            "apertura local comercial permiso funcionamiento actividad economica"
        ),
        "priority": "IMPORTANTE",
        "priority_level": 2,
        "office": "Gerencia de Desarrollo Económico",
        "office_icon": "🏪",
        "base_score": 50,
        "type_bonus": 5,
        "color": "importante",
    },
    "Partición o Subdivisión de Lote": {
        "keywords": (
            "lote particion subdivision catastro predio desmembracion terreno propiedad "
            "registro area metraje colindantes lote matriz sublote fraccionar lotizacion "
            "plano division parcela hectarea linderos superficie"
        ),
        "priority": "IMPORTANTE",
        "priority_level": 2,
        "office": "Gerencia de Urbanismo y Obras",
        "office_icon": "🏗️",
        "base_score": 50,
        "type_bonus": 4,
        "color": "importante",
    },

    # ── RUTINARIO (Baja) ──────────────────────────────────────────────────────
    "Solicitud de Información Pública": {
        "keywords": (
            "informacion consulta solicito acceso transparencia datos publicos informe reporte "
            "pedido conocer estado tramite ley transparencia acceso informacion solicitud copia "
            "expediente informacion publica ciudadano derecho"
        ),
        "priority": "RUTINARIO",
        "priority_level": 1,
        "office": "Mesa de Partes General",
        "office_icon": "📋",
        "base_score": 20,
        "type_bonus": 2,
        "color": "rutinario",
    },
    "Pago de Arbitrios e Impuestos": {
        "keywords": (
            "arbitrios impuesto predial tributo pago contribuyente deuda municipal recibo "
            "cuota anual limpieza parques alumbrado serenazgo declaracion jurada autovaluo "
            "fraccionamiento deuda tributaria impuesto municipal"
        ),
        "priority": "RUTINARIO",
        "priority_level": 1,
        "office": "Gerencia de Rentas",
        "office_icon": "💰",
        "base_score": 20,
        "type_bonus": 2,
        "color": "rutinario",
    },
    "Certificado de Residencia": {
        "keywords": (
            "residencia domicilio vecino habita certificado morador direccion vivir vivienda "
            "habitante padron domicilio declarado certifica reside constancia domiciliario "
            "ficha padron vecinal barrio distrito"
        ),
        "priority": "RUTINARIO",
        "priority_level": 1,
        "office": "Mesa de Partes General",
        "office_icon": "📋",
        "base_score": 20,
        "type_bonus": 3,
        "color": "rutinario",
    },
    "Constancia de Posesión": {
        "keywords": (
            "posesion terreno ocupante constancia poseedor parcela campo agricola rural "
            "zona posesionario posesion pacifica publica continua propietario titular predio "
            "rustico urbano posesion directa terreno baldio"
        ),
        "priority": "RUTINARIO",
        "priority_level": 1,
        "office": "Gerencia de Catastro",
        "office_icon": "🗺️",
        "base_score": 20,
        "type_bonus": 3,
        "color": "rutinario",
    },
    "Queja o Reclamo Administrativo": {
        "keywords": (
            "queja reclamo inconformidad insatisfaccion protesta servicio mal atencion "
            "deficiente problema funcionario negligencia demora incumplimiento queja ciudadana "
            "libro reclamaciones mala atencion insatisfecho disconforme"
        ),
        "priority": "RUTINARIO",
        "priority_level": 1,
        "office": "Mesa de Partes General",
        "office_icon": "📋",
        "base_score": 20,
        "type_bonus": 2,
        "color": "rutinario",
    },
}

TRAMITE_NAMES = list(TRAMITE_CATALOG.keys())
TRAMITE_PROFILES = list(TRAMITE_CATALOG.values())


# ─── TEXT EXTRACTION ──────────────────────────────────────────────────────────

def _extract_pdf(content: bytes) -> str:
    import pdfplumber
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        parts = []
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                parts.append(t)
        return "\n".join(parts)


def _extract_docx(content: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(content))
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    table_rows = []
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text for c in row.cells if c.text.strip())
            if row_text:
                table_rows.append(row_text)
    return "\n".join(paras + table_rows)


def extract_text(content: bytes, filename: str) -> Tuple[str, Optional[str]]:
    """Extract text from PDF or DOCX. Returns (text, error_message)."""
    try:
        lower = filename.lower()
        if lower.endswith(".pdf"):
            text = _extract_pdf(content)
        elif lower.endswith(".docx"):
            text = _extract_docx(content)
        else:
            return "", "Formato no soportado. Use PDF o DOCX."

        if not text.strip():
            return "", (
                "No se pudo extraer texto del documento. "
                "El archivo puede ser un escaneado (imagen sin texto seleccionable) "
                "o estar vacío."
            )
        return text.strip(), None
    except Exception as exc:
        return "", f"Error al procesar el archivo: {str(exc)}"


# ─── AGE BONUS ────────────────────────────────────────────────────────────────

def _extract_document_year(text: str) -> Optional[int]:
    """Extract the earliest plausible year from document text."""
    current_year = datetime.now().year

    # DD/MM/YYYY or YYYY-MM-DD structured dates
    structured = re.findall(
        r'\b(?:\d{1,2}[/\-]\d{1,2}[/\-]((?:19|20)\d{2})|((?:19|20)\d{2})[/\-]\d{1,2}[/\-]\d{1,2})\b',
        text,
    )
    struct_years = [int(y) for pair in structured for y in pair if y]

    # Any bare year
    bare = re.findall(r'\b((?:19[5-9]|20[0-2])\d)\b', text)
    bare_years = [int(y) for y in bare]

    candidates = struct_years + bare_years
    valid = [y for y in candidates if 1950 <= y <= current_year]
    return min(valid) if valid else None


def calculate_age_bonus(text: str) -> Tuple[Optional[int], int, str]:
    """Returns (document_year, bonus_points, description)."""
    doc_year = _extract_document_year(text)
    if doc_year is None:
        return None, 0, "Sin fecha detectada en el documento"

    age = datetime.now().year - doc_year
    if age > 5:
        return doc_year, 20, f"Documento del año {doc_year} — más de 5 años de antigüedad (+20 pts)"
    elif age >= 2:
        return doc_year, 12, f"Documento del año {doc_year} — {age} años de antigüedad (+12 pts)"
    elif age >= 1:
        return doc_year, 6, f"Documento del año {doc_year} — 1 año de antigüedad (+6 pts)"
    else:
        return doc_year, 0, f"Documento reciente (año {doc_year}) — sin bonus de antigüedad"


# ─── NORMALIZATION ────────────────────────────────────────────────────────────

_ACCENT_MAP = str.maketrans(
    "áéíóúüñÁÉÍÓÚÜÑ",
    "aeiouunAEIOUUN",
)


def _normalize(text: str) -> str:
    return text.lower().translate(_ACCENT_MAP)


# ─── CLASSIFIER ───────────────────────────────────────────────────────────────

def classify_document(text: str) -> dict:
    """Clasifica el texto de un documento usando TF-IDF y similitud coseno."""
    norm_text = _normalize(text)
    norm_keywords = [_normalize(p["keywords"]) for p in TRAMITE_PROFILES]

    corpus = norm_keywords + [norm_text]

    try:
        vectorizer = TfidfVectorizer(
            analyzer="word",
            ngram_range=(1, 2),
            min_df=1,
            sublinear_tf=True,
            max_features=8000,
        )
        tfidf = vectorizer.fit_transform(corpus)
        similarities = cosine_similarity(tfidf[-1:], tfidf[:-1])[0]
    except Exception:
        # Fallback
        return _keyword_fallback(norm_text, text)

    best_idx = int(np.argmax(similarities))
    confidence = float(similarities[best_idx])
    best_name = TRAMITE_NAMES[best_idx]
    profile = TRAMITE_PROFILES[best_idx]

    doc_year, age_bonus, age_desc = calculate_age_bonus(text)

    final_score = min(100, profile["base_score"] + profile["type_bonus"] + age_bonus)

    # Alternativas
    top5_idx = np.argsort(similarities)[::-1][:5]
    alternatives = [
        {
            "name": TRAMITE_NAMES[i],
            "priority": TRAMITE_PROFILES[i]["priority"],
            "similarity": round(float(similarities[i]) * 100, 1),
        }
        for i in top5_idx
    ]

    return {
        "tramite_type": best_name,
        "priority": profile["priority"],
        "priority_level": profile["priority_level"],
        "office": profile["office"],
        "office_icon": profile["office_icon"],
        "score": final_score,
        "score_breakdown": {
            "base": profile["base_score"],
            "type_bonus": profile["type_bonus"],
            "age_bonus": age_bonus,
            "total": final_score,
        },
        "confidence": round(confidence * 100, 1),
        "age_bonus": age_bonus,
        "age_description": age_desc,
        "document_year": doc_year,
        "text_excerpt": text[:400].strip(),
        "color": profile["color"],
        "alternatives": alternatives,
    }


def _keyword_fallback(norm_text: str, raw_text: str) -> dict:
    """Búsqueda simple por palabras clave en caso de fallback."""
    scores = []
    for i, profile in enumerate(TRAMITE_PROFILES):
        kws = _normalize(profile["keywords"]).split()
        count = sum(1 for kw in kws if kw in norm_text)
        scores.append(count)

    best_idx = int(np.argmax(scores)) if scores else 0
    best_name = TRAMITE_NAMES[best_idx]
    profile = TRAMITE_PROFILES[best_idx]
    doc_year, age_bonus, age_desc = calculate_age_bonus(raw_text)
    final_score = min(100, profile["base_score"] + profile["type_bonus"] + age_bonus)

    return {
        "tramite_type": best_name,
        "priority": profile["priority"],
        "priority_level": profile["priority_level"],
        "office": profile["office"],
        "office_icon": profile["office_icon"],
        "score": final_score,
        "score_breakdown": {
            "base": profile["base_score"],
            "type_bonus": profile["type_bonus"],
            "age_bonus": age_bonus,
            "total": final_score,
        },
        "confidence": round(min(scores[best_idx] * 8.0, 100.0), 1),
        "age_bonus": age_bonus,
        "age_description": age_desc,
        "document_year": doc_year,
        "text_excerpt": raw_text[:400].strip(),
        "color": profile["color"],
        "alternatives": [],
    }
